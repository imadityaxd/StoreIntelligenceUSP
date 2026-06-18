from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "Computer_Vision_Tools_OpenCV_YOLO_Tracking_Guide.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "5E6A78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "B7C3D0"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: str = "6") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.autofit = False


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    page_break = OxmlElement("w:pageBreakBefore")
    p_pr.append(page_break)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Vivid Store AI reference guide")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Computer Vision Building Blocks")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("OpenCV, YOLOv8 / Ultralytics, ByteTrack, and BoT-SORT explained for Vivid Store AI")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    add_callout(
        doc,
        "Purpose of this guide",
        "This document explains the three main computer vision layers in the CCTV analytics pipeline. "
        "It answers the four W's for each tool: what it is, why it matters, where it is used, and when it runs. "
        "The goal is to help you understand the technology behind person boxes, IDs, and live store metrics.",
    )


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    set_table_borders(table, color="D5DCE5", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r.font.size = Pt(11)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_four_ws_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table)
    set_table_borders(table)
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.15)
    header = table.rows[0]
    set_repeat_table_header(header)
    header.cells[0].text = "Question"
    header.cells[1].text = "Answer"
    for cell in header.cells:
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for question, answer in rows:
        cells = table.add_row().cells
        cells[0].text = question
        cells[1].text = answer
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()


def add_comparison_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table)
    set_table_borders(table)
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()


def add_intro(doc: Document) -> None:
    doc.add_heading("1. The Simple Pipeline", level=1)
    p = doc.add_paragraph()
    p.add_run("In your project, these tools do not work separately. ").bold = True
    p.add_run(
        "They form a pipeline: OpenCV reads CCTV frames, YOLOv8 finds people in those frames, "
        "ByteTrack or BoT-SORT links detections across time, and the application turns stable tracks into overlays and metrics."
    )
    add_comparison_table(
        doc,
        ["Stage", "Tool", "Main job", "Output"],
        [
            ["1", "OpenCV", "Open video or camera stream and extract frames.", "Image frames with width, height, FPS, and timestamp context."],
            ["2", "YOLOv8 / Ultralytics", "Detect people in each sampled frame.", "Person bounding boxes with confidence scores."],
            ["3", "ByteTrack / BoT-SORT", "Connect boxes across frames into identities.", "Track IDs that stay stable while a person remains visible."],
            ["4", "Vivid Store AI pipeline", "Validate physical people, map zones, generate events, and draw overlays.", "Green countable people, suspect boxes, events, metrics, and reports."],
        ],
        widths=[0.55, 1.55, 2.9, 2.0],
    )
    add_callout(
        doc,
        "Memory hook",
        "OpenCV handles frames. YOLO handles detection. ByteTrack and BoT-SORT handle identity over time. "
        "The app handles business meaning.",
    )


def add_opencv(doc: Document) -> None:
    doc.add_heading("2. OpenCV", level=1)
    add_four_ws_table(
        doc,
        [
            ("What?", "OpenCV is a computer vision library used to read, write, transform, and analyze images and video frames. In this project, it is the frame and pixel handling layer."),
            ("Why?", "YOLO cannot work directly on a video file as a business object. It needs image frames. OpenCV opens the CCTV clip or stream, reads frames, reports FPS and resolution, and gives the detector actual pixel arrays."),
            ("Where?", "OpenCV is used in the detection pipeline, video inspection, frame extraction, camera calibration helpers, and any backend-side overlay or reference frame generation."),
            ("When?", "It runs before YOLO for frame extraction, during processing for frame sampling and geometry, and after detection when frames or overlays need to be saved, inspected, or exported."),
        ],
    )
    doc.add_heading("Core Ideas", level=2)
    add_bullets(
        doc,
        [
            "Frame: one still image from the video. A 30 FPS video has about 30 frames per second.",
            "FPS: frames per second. This controls timestamp math and how often the system samples the video.",
            "Resolution: width and height of the frame, such as 1920 x 1080. Detection boxes use these coordinates.",
            "Color format: OpenCV usually reads images as BGR, while many AI/image tools think in RGB.",
            "VideoCapture: the OpenCV object that opens a file, webcam, or stream and reads frames one by one.",
        ],
    )
    doc.add_heading("What OpenCV Does Not Do Here", level=2)
    add_bullets(
        doc,
        [
            "It does not understand who is a person by itself in the current YOLO-based pipeline.",
            "It does not maintain identity across frames by itself.",
            "It does not decide whether someone entered a zone, joined a queue, or converted.",
        ],
    )
    add_callout(
        doc,
        "In one line",
        "OpenCV is the video and pixel plumbing that feeds the AI model clean frames.",
    )


def add_yolo(doc: Document) -> None:
    doc.add_heading("3. YOLOv8 / Ultralytics", level=1)
    add_four_ws_table(
        doc,
        [
            ("What?", "YOLOv8 is an object detection model. Ultralytics is the Python package/framework commonly used to run YOLOv8 models. It returns detected objects, including person boxes, class labels, and confidence scores."),
            ("Why?", "The dashboard needs to know where people are in each frame. YOLOv8 provides person bounding boxes more reliably than simple motion detection, especially when the camera is fixed but lighting, shelves, and background are complex."),
            ("Where?", "YOLOv8 runs inside the detection pipeline after OpenCV extracts a sampled frame. Its output is passed to the tracker and also becomes the basis for the visual boxes shown on the dashboard."),
            ("When?", "It runs on each sampled frame during a CCTV analysis session or live stream session. Sampling is controlled by process FPS so the system can balance speed and accuracy."),
        ],
    )
    doc.add_heading("Core Ideas", level=2)
    add_bullets(
        doc,
        [
            "Object detection: finding where an object is, not just saying what is in the image.",
            "Bounding box: the rectangle around a detected person, usually represented as x1, y1, x2, y2.",
            "Class: the object category, such as person, bag, chair, or bottle. This project mainly uses the person class.",
            "Confidence: the model's score for how likely the box is a real object of that class.",
            "IoU: Intersection over Union. It measures how much two boxes overlap and is used during filtering and tracking.",
            "NMS: Non-Max Suppression. It removes duplicate boxes around the same object.",
            "Image size: the input size used for inference. Larger sizes can detect smaller people but cost more CPU/GPU time.",
        ],
    )
    doc.add_heading("Why YOLO Alone Is Not Enough", level=2)
    add_bullets(
        doc,
        [
            "YOLO detects a person in one frame, but it does not automatically know that the same person appears in the next frame.",
            "YOLO can detect mirror reflections, posters, partial bodies, or weak shapes as person-like boxes.",
            "YOLO does not know retail business meaning such as queue, dwell, funnel, or conversion.",
        ],
    )
    add_callout(
        doc,
        "In one line",
        "YOLOv8 is the model that says: there is a person-shaped object at this rectangle in this frame.",
    )


def add_tracking(doc: Document) -> None:
    doc.add_heading("4. ByteTrack and BoT-SORT", level=1)
    add_four_ws_table(
        doc,
        [
            ("What?", "ByteTrack and BoT-SORT are multi-object tracking algorithms. They connect YOLO detections across frames and assign track IDs so one person can keep the same identity while visible."),
            ("Why?", "Without tracking, the system would see the same person as a new detection in every frame. Tracking prevents repeated counting and makes movement, dwell time, queue behavior, and event generation possible."),
            ("Where?", "They run after YOLO detection and before event generation. The tracker output feeds the overlay labels, physical-person validation, zone mapping, and business metrics."),
            ("When?", "They run continuously during analysis. Every sampled frame updates existing tracks, creates new tracks, or marks missing tracks as temporarily lost or ended."),
        ],
    )
    doc.add_heading("Core Ideas", level=2)
    add_bullets(
        doc,
        [
            "Track ID: a temporary identity assigned to a detected person inside one camera run.",
            "Association: the process of deciding which new box belongs to which existing track.",
            "Motion prediction: estimating where a track should appear in the next frame.",
            "Lost track: a person was visible before but is missing now, maybe due to occlusion or detection failure.",
            "Re-identification: using appearance features to reconnect an identity after a short disappearance. BoT-SORT can use stronger appearance cues depending on setup.",
        ],
    )
    doc.add_heading("ByteTrack vs BoT-SORT", level=2)
    add_comparison_table(
        doc,
        ["Tracker", "Strength", "Best use", "Tradeoff"],
        [
            ["ByteTrack", "Uses high-confidence and low-confidence detections to keep tracks alive.", "Good general-purpose tracking when people are visible and movement is moderate.", "Can still switch IDs when occlusion is heavy or camera views are complex."],
            ["BoT-SORT", "Combines motion, box matching, and stronger association logic; can be better in crowded or difficult scenes.", "Billing counters, dense zones, and camera views with partial occlusions.", "Usually heavier and may need more tuning than simpler tracking."],
        ],
        widths=[1.1, 2.1, 2.05, 1.25],
    )
    add_callout(
        doc,
        "In one line",
        "Tracking is what turns many frame-by-frame detections into one continuous person path.",
    )


def add_together(doc: Document) -> None:
    doc.add_heading("5. How They Work Together In Vivid Store AI", level=1)
    add_bullets(
        doc,
        [
            "OpenCV opens the CCTV file or stream and extracts frames.",
            "The pipeline samples frames based on the selected process FPS.",
            "YOLOv8 detects people in each sampled frame.",
            "ByteTrack or BoT-SORT links person boxes across frames into track IDs.",
            "Physical-person validation checks whether a track is stable and plausible enough to count.",
            "The zone system maps each valid person's footpoint to store zones.",
            "The event emitter creates entries, zone visits, queue joins, exits, and dwell events when supported.",
            "The dashboard draws boxes, IDs, metrics, event feed, funnel, heatmap, and reports.",
        ],
    )
    add_comparison_table(
        doc,
        ["Problem you see", "Likely layer", "What to tune or inspect"],
        [
            ["No boxes on people", "YOLO/OpenCV", "Check video readability, lighting, YOLO confidence, image size, and process FPS."],
            ["Boxes appear but IDs change", "Tracker", "Try Auto profile, ByteTrack, BoT-SORT, or higher process FPS."],
            ["Mirror/reflection gets counted", "Validation layer", "Review suspect boxes, confidence, geometry, and physical-person validation strictness."],
            ["Boxes work but zones are empty", "Layout/calibration", "Check camera role, store layout polygons, and zone mapping."],
            ["Metrics are zero after a run", "Events/API/session", "Check whether valid tracks generated events and whether dashboard is scoped to the current session."],
        ],
        widths=[1.7, 1.3, 3.5],
    )


def add_glossary(doc: Document) -> None:
    doc.add_heading("6. Glossary", level=1)
    add_comparison_table(
        doc,
        ["Term", "Meaning"],
        [
            ["Frame", "A single image from a video."],
            ["Detection", "A model output saying an object exists at a bounding box."],
            ["Bounding box", "The rectangle around a detected person."],
            ["Track", "A sequence of detections linked as the same person over time."],
            ["Track ID", "The temporary identity assigned to one tracked person."],
            ["Confidence", "The model's score for a detection."],
            ["IoU", "A measure of how much two boxes overlap."],
            ["Occlusion", "When a person is partly or fully hidden by another object or person."],
            ["Dwell", "How long a person stays in a zone."],
            ["Footpoint", "The bottom-center point of a box, often used to estimate where a person stands on the floor."],
        ],
        widths=[1.55, 4.95],
    )


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_title(doc)
    add_intro(doc)
    add_opencv(doc)
    add_yolo(doc)
    add_tracking(doc)
    add_together(doc)
    add_glossary(doc)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
